"""
Generate two demo .docx files to test the deployed converter:

  Demo_MCQ_with_images.docx   4 questions, each with an embedded image
  Demo_MCQ_no_images.docx     6 questions, plain text only

Each file covers all three correct-answer detection styles:
  1. "Answer: X" text marker
  2. Highlighted correct option
  3. Bold correct option
"""
import os

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches


HERE = os.path.dirname(os.path.abspath(__file__))
TMP = os.path.join(HERE, "_demo_images")
os.makedirs(TMP, exist_ok=True)


def _font(size=48):
    for f in ("arial.ttf", "Arial.ttf", "DejaVuSans-Bold.ttf"):
        try:
            return ImageFont.truetype(f, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_square(colour, path):
    img = Image.new("RGB", (200, 200), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([30, 30, 170, 170], fill=colour, outline="black", width=3)
    img.save(path)


def draw_circle(path):
    img = Image.new("RGB", (240, 200), "white")
    d = ImageDraw.Draw(img)
    d.ellipse([60, 30, 180, 170], outline="black", width=4, fill="#87ceeb")
    img.save(path)


def draw_dots(n, path):
    img = Image.new("RGB", (240, 120), "white")
    d = ImageDraw.Draw(img)
    for i in range(n):
        cx = 40 + i * 55
        d.ellipse([cx - 18, 42, cx + 18, 78], fill="black")
    img.save(path)


def draw_letter(ch, path):
    img = Image.new("RGB", (200, 200), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([4, 4, 195, 195], outline="black", width=3)
    font = _font(140)
    bbox = d.textbbox((0, 0), ch, font=font)
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    d.text(((200 - w) / 2 - bbox[0], (200 - h) / 2 - bbox[1]), ch,
           fill="#c00020", font=font)
    img.save(path)


# ─── File 1: with images ─────────────────────────────────────────────────────

def build_with_images():
    img1 = os.path.join(TMP, "red_square.png")
    img2 = os.path.join(TMP, "circle.png")
    img3 = os.path.join(TMP, "three_dots.png")
    img4 = os.path.join(TMP, "letter_A.png")
    draw_square("#d62828", img1)
    draw_circle(img2)
    draw_dots(3, img3)
    draw_letter("A", img4)

    doc = Document()
    doc.add_heading("Demo MCQ — with images", level=1)

    # Q1: Answer: X
    doc.add_paragraph("Q1")
    doc.add_paragraph("What colour is the square shown below?")
    doc.add_paragraph().add_run().add_picture(img1, width=Inches(1.5))
    doc.add_paragraph("A) Blue")
    doc.add_paragraph("B) Red")
    doc.add_paragraph("C) Green")
    doc.add_paragraph("D) Yellow")
    doc.add_paragraph("Answer: B")
    doc.add_paragraph("Explanation: the square is clearly red.")

    # Q2: highlighted correct option
    doc.add_paragraph("Q2")
    doc.add_paragraph("Which shape is shown in the image?")
    doc.add_paragraph().add_run().add_picture(img2, width=Inches(1.8))
    doc.add_paragraph("A) Square")
    p = doc.add_paragraph()
    r = p.add_run("B) Circle")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_paragraph("C) Triangle")
    doc.add_paragraph("D) Pentagon")

    # Q3: bold correct option
    doc.add_paragraph("Q3")
    doc.add_paragraph("How many dots are in the image below?")
    doc.add_paragraph().add_run().add_picture(img3, width=Inches(2.0))
    doc.add_paragraph("A) 2")
    p = doc.add_paragraph()
    p.add_run("B) 3").bold = True
    doc.add_paragraph("C) 4")
    doc.add_paragraph("D) 5")

    # Q4: Answer: X
    doc.add_paragraph("Q4")
    doc.add_paragraph("Which letter is displayed?")
    doc.add_paragraph().add_run().add_picture(img4, width=Inches(1.5))
    doc.add_paragraph("A) A")
    doc.add_paragraph("B) B")
    doc.add_paragraph("C) C")
    doc.add_paragraph("D) D")
    doc.add_paragraph("Answer: A")

    out = os.path.join(HERE, "Demo_MCQ_with_images.docx")
    doc.save(out)
    return out


# ─── File 2: no images ───────────────────────────────────────────────────────

def build_no_images():
    doc = Document()
    doc.add_heading("Demo MCQ — text only", level=1)

    # Q1: Answer: X
    doc.add_paragraph("Q1")
    doc.add_paragraph("What is the capital of France?")
    doc.add_paragraph("A) Berlin")
    doc.add_paragraph("B) Madrid")
    doc.add_paragraph("C) Paris")
    doc.add_paragraph("D) Rome")
    doc.add_paragraph("Answer: C")

    # Q2: highlight
    doc.add_paragraph("Q2")
    doc.add_paragraph("Which planet is known as the Red Planet?")
    doc.add_paragraph("A) Venus")
    doc.add_paragraph("B) Jupiter")
    p = doc.add_paragraph()
    r = p.add_run("C) Mars")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_paragraph("D) Saturn")

    # Q3: bold
    doc.add_paragraph("Q3")
    doc.add_paragraph("How many continents are there on Earth?")
    doc.add_paragraph("A) 5")
    doc.add_paragraph("B) 6")
    p = doc.add_paragraph()
    p.add_run("C) 7").bold = True
    doc.add_paragraph("D) 8")
    doc.add_paragraph("Explanation: Africa, Antarctica, Asia, Australia, Europe, North & South America.")

    # Q4: Answer: X
    doc.add_paragraph("Q4")
    doc.add_paragraph("What is 12 * 8?")
    doc.add_paragraph("A) 86")
    doc.add_paragraph("B) 96")
    doc.add_paragraph("C) 104")
    doc.add_paragraph("D) 112")
    doc.add_paragraph("Answer: B")

    # Q5: highlight
    doc.add_paragraph("Q5")
    doc.add_paragraph("Which language is this app's backend written in?")
    p = doc.add_paragraph()
    r = p.add_run("A) Python")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_paragraph("B) Rust")
    doc.add_paragraph("C) Go")
    doc.add_paragraph("D) Ruby")

    # Q6: bold
    doc.add_paragraph("Q6")
    doc.add_paragraph("Which HTTP status code means 'Not Found'?")
    doc.add_paragraph("A) 200")
    doc.add_paragraph("B) 301")
    doc.add_paragraph("C) 403")
    p = doc.add_paragraph()
    p.add_run("D) 404").bold = True

    out = os.path.join(HERE, "Demo_MCQ_no_images.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    a = build_with_images()
    b = build_no_images()
    for p in (a, b):
        size = os.path.getsize(p)
        print(f"  {os.path.basename(p):35s} {size:>8} bytes")
